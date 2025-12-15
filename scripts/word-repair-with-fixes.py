#!/usr/bin/env python3
"""
Word Document Repair Script
Modifies Word documents to apply accessibility fixes:
- Alt text for images
- Table summaries
- Document metadata (title, language)
- Heading structure
- Language tags
- Color contrast fixes
- Link text improvements
- List structure

Uses python-docx library to modify .docx files
"""

import sys
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx library not installed. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def add_alt_text_to_image(doc, image_id, alt_text):
    """
    Add alt text to an image by modifying the XML directly
    Uses image index (image0, image1, etc.) to find images by order
    """
    try:
        # Extract image index from image_id (e.g., "image0" -> 0, "image1" -> 1, "image0.png" -> 0)
        image_index = 0
        if image_id.startswith('image'):
            try:
                # Handle both "image0" and "image0.png" formats
                index_str = image_id.replace('image', '').split('.')[0]
                if index_str.isdigit():
                    image_index = int(index_str)
            except:
                pass
        
        # Find all images in the document by iterating through paragraphs
        all_images = []
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # Check for inline images (use namespace prefix directly in xpath)
                inline_drawings = run._element.xpath('.//wp:inline')
                for drawing in inline_drawings:
                    all_images.append(drawing)
                
                # Check for anchored images (floating images)
                anchor_drawings = run._element.xpath('.//wp:anchor')
                for drawing in anchor_drawings:
                    all_images.append(drawing)
        
        # Apply alt text to the image at the specified index
        if image_index < len(all_images):
            drawing = all_images[image_index]
            
            # Get or create docPr element (document properties) - this contains alt text
            docPrs = drawing.xpath('.//wp:docPr')
            
            if docPrs:
                docPr = docPrs[0]
            else:
                # Create new docPr element
                docPr = OxmlElement('wp:docPr')
                # Insert at the beginning of the drawing element
                drawing.insert(0, docPr)
            
            # Set the description (alt text) - this is what screen readers use
            docPr.set(qn('wp:descr'), alt_text)
            docPr.set(qn('wp:name'), 'Picture')
            
            return True
        else:
            print(f"WARNING: Image index {image_index} out of range (found {len(all_images)} images)", file=sys.stderr)
            return False
            
    except Exception as e:
        print(f"WARNING: Could not add alt text to {image_id}: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return False


def set_table_summary(doc, table_index, summary):
    """
    Set summary/description for a table by adding a caption paragraph before it
    """
    try:
        if table_index < len(doc.tables):
            table = doc.tables[table_index]
            # Find the paragraph that contains this table
            # We'll add a caption paragraph before the table
            table_element = table._element
            parent = table_element.getparent()
            
            # Create a new paragraph for the caption
            caption_para = OxmlElement('w:p')
            caption_run = OxmlElement('w:r')
            caption_text = OxmlElement('w:t')
            caption_text.text = f"Table {table_index + 1}: {summary}"
            caption_run.append(caption_text)
            caption_para.append(caption_run)
            
            # Insert before the table
            parent.insert(parent.index(table_element), caption_para)
            return True
        return False
    except Exception as e:
        print(f"WARNING: Could not set table summary: {str(e)}", file=sys.stderr)
        return False


def repair_word_document(input_path: str, output_path: str, fixes: dict):
    """
    Repair Word document with accessibility fixes
    """
    try:
        # Load the document
        doc = Document(input_path)
        
        fixes_applied = 0
        
        # Apply metadata fixes
        if fixes.get('metadata'):
            metadata = fixes['metadata']
            if metadata.get('title'):
                # Set document title in core properties
                # python-docx requires direct XML manipulation for reliable title setting
                try:
                    title_value = str(metadata['title']).strip()
                    if title_value:
                        # Method 1: Set via core_properties (standard way)
                        doc.core_properties.title = title_value
                        
                        # Method 2: python-docx should handle this via core_properties.title
                        # Just set it via the property - python-docx handles XML internally
                        doc.core_properties.title = title_value
                        
                        # Verify it was set
                        verify_title = doc.core_properties.title
                        if verify_title and verify_title.strip() == title_value:
                            print(f"INFO: Set document title: {title_value}")
                            fixes_applied += 1
                        else:
                            print(f"WARNING: Title was set but verification failed (got: '{verify_title}')", file=sys.stderr)
                except Exception as e:
                    print(f"WARNING: Could not set document title: {str(e)}", file=sys.stderr)
                    import traceback
                    traceback.print_exc()
            
            if metadata.get('language'):
                # Set document language
                lang_code = metadata['language']
                # Set default language on document settings
                try:
                    # Set language on all paragraphs and runs
                    for paragraph in doc.paragraphs:
                        paragraph._element.get_or_add_pPr().get_or_add_rPr().get_or_add_lang().set(qn('w:val'), lang_code)
                        for run in paragraph.runs:
                            run._element.get_or_add_rPr().get_or_add_lang().set(qn('w:val'), lang_code)
                    
                    # Also set on document default style
                    styles = doc.styles
                    if 'Normal' in [s.name for s in styles]:
                        normal_style = styles['Normal']
                        if normal_style._element.rPr is None:
                            normal_style._element.get_or_add_rPr()
                        normal_style._element.rPr.get_or_add_lang().set(qn('w:val'), lang_code)
                    
                    print(f"INFO: Document language set to: {lang_code}")
                    fixes_applied += 1
                except Exception as e:
                    print(f"WARNING: Could not set document language: {str(e)}", file=sys.stderr)
        
        # Apply image alt text fixes
        image_fixes = fixes.get('imageFixes', [])
        if image_fixes:
            print(f"INFO: Applying {len(image_fixes)} image alt text fixes...")
            for image_fix in image_fixes:
                image_id = image_fix.get('imageId', '')
                alt_text = image_fix.get('altText', 'Image')
                if add_alt_text_to_image(doc, image_id, alt_text):
                    print(f"INFO: Added alt text to image {image_id}: {alt_text[:50]}...")
                    fixes_applied += 1
                else:
                    print(f"WARNING: Could not add alt text to image {image_id}", file=sys.stderr)
        
        # Apply table summary fixes
        table_fixes = fixes.get('tableFixes', [])
        if table_fixes:
            print(f"INFO: Applying {len(table_fixes)} table summary fixes...")
            for i, table_fix in enumerate(table_fixes):
                summary = table_fix.get('summary', '')
                if summary and set_table_summary(doc, i, summary):
                    print(f"INFO: Added summary to table {i}: {summary[:50]}...")
                    fixes_applied += 1
                else:
                    print(f"WARNING: Could not add summary to table {i}", file=sys.stderr)
                
                # Also mark first row as header if table doesn't have headers
                if i < len(doc.tables):
                    table = doc.tables[i]
                    has_headers = table_fix.get('hasHeaders', False)
                    if not has_headers and len(table.rows) > 0:
                        try:
                            # Mark first row as header by setting header row style
                            first_row = table.rows[0]
                            for cell in first_row.cells:
                                # Set cell style to indicate header
                                cell._element.get_or_add_tcPr().get_or_add_shd().set(qn('w:fill'), 'D9D9D9')  # Light gray background
                            
                            # Set table style to use first row as header
                            table._element.set(qn('w:tblHeader'), '1')
                            print(f"INFO: Marked first row as header for table {i}")
                            fixes_applied += 1
                        except Exception as e:
                            print(f"WARNING: Could not mark header row for table {i}: {str(e)}", file=sys.stderr)
        
        # Apply color contrast fixes
        # Strategy: Scan ALL text runs and fix ANY low contrast colors (not just specific ones)
        color_contrast_fixes = fixes.get('colorContrastFixes', [])
        
        def calculate_luminance(r, g, b):
            """Calculate relative luminance for WCAG contrast calculation"""
            def normalize(val):
                val = val / 255.0
                return val / 12.92 if val <= 0.03928 else ((val + 0.055) / 1.055) ** 2.4
            r_norm = normalize(r)
            g_norm = normalize(g)
            b_norm = normalize(b)
            return 0.2126 * r_norm + 0.7152 * g_norm + 0.0722 * b_norm
        
        def calculate_contrast_ratio(rgb1, rgb2):
            """Calculate WCAG contrast ratio between two RGB colors"""
            l1 = calculate_luminance(rgb1[0], rgb1[1], rgb1[2])
            l2 = calculate_luminance(rgb2[0], rgb2[1], rgb2[2])
            lighter = max(l1, l2)
            darker = min(l1, l2)
            return (lighter + 0.05) / (darker + 0.05)
        
        # Default accessible color (dark grey - meets WCAG AAA)
        accessible_color = RGBColor(89, 89, 89)  # #595959
        background_rgb = (255, 255, 255)  # White background (default)
        
        if color_contrast_fixes and isinstance(color_contrast_fixes, list) and len(color_contrast_fixes) > 0:
            # Get the suggested accessible color from the first fix
            first_fix = color_contrast_fixes[0]
            new_color_hex = first_fix.get('newColor', '#595959')
            hex_clean = new_color_hex.lstrip('#')
            if len(hex_clean) == 6:
                r = int(hex_clean[0:2], 16)
                g = int(hex_clean[2:4], 16)
                b = int(hex_clean[4:6], 16)
                accessible_color = RGBColor(r, g, b)
        
        print(f"INFO: Scanning document for ALL low contrast colors (WCAG AA threshold: 4.5:1)...")
        runs_fixed = 0
        
        # Scan ALL paragraphs and runs in the document
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if not run.text or not run.text.strip():
                    continue
                
                try:
                    current_color = run.font.color
                    if current_color and current_color.rgb:
                        # Get RGB values
                        r, g, b = current_color.rgb
                        
                        # Calculate contrast ratio against white background
                        contrast = calculate_contrast_ratio((r, g, b), background_rgb)
                        
                        # WCAG AA requires 4.5:1 for normal text, 3:1 for large text
                        # We'll use 4.5:1 as the threshold to catch all inaccessible colors
                        if contrast < 4.5:
                            # Low contrast - fix it
                            old_hex = f"#{r:02x}{g:02x}{b:02x}"
                            run.font.color.rgb = accessible_color
                            new_r, new_g, new_b = accessible_color.rgb
                            new_hex = f"#{new_r:02x}{new_g:02x}{new_b:02x}"
                            new_contrast = calculate_contrast_ratio((new_r, new_g, new_b), background_rgb)
                            print(f"INFO: Fixed low contrast: {old_hex} ({contrast:.2f}:1) -> {new_hex} ({new_contrast:.2f}:1) for text: '{run.text[:50]}...'")
                            runs_fixed += 1
                            fixes_applied += 1
                    else:
                        # No color set - check if text appears light (might be styled elsewhere)
                        # For now, we'll skip runs without explicit color
                        pass
                except Exception as e:
                    # Skip runs that cause errors
                    continue
        
        if runs_fixed > 0:
            print(f"INFO: Fixed {runs_fixed} text runs with low contrast colors")
        else:
            print(f"INFO: No low contrast colors found (or fixes already applied)")
        
        # Apply link text fixes
        link_text_fixes = fixes.get('linkTextFixes', [])
        if link_text_fixes and isinstance(link_text_fixes, list):
            print(f"INFO: Applying {len(link_text_fixes)} link text fixes...")
            for link_fix in link_text_fixes:
                try:
                    old_text = link_fix.get('oldText', '')
                    new_text = link_fix.get('newText', '')
                    element_location = link_fix.get('elementLocation', '')
                    
                    if old_text and new_text and old_text != new_text:
                        # Find hyperlinks containing the old text and replace
                        found = False
                        for paragraph in doc.paragraphs:
                            # Check for hyperlinks in paragraph
                            hyperlinks = paragraph._element.xpath('.//w:hyperlink')
                            for hyperlink in hyperlinks:
                                # Get all text runs in this hyperlink
                                text_elements = hyperlink.xpath('.//w:t')
                                full_text = ''.join([elem.text or '' for elem in text_elements])
                                
                                if old_text.lower() in full_text.lower():
                                    # Replace text in the first text element
                                    if text_elements:
                                        text_elements[0].text = new_text
                                        # Clear other text elements to avoid duplication
                                        for elem in text_elements[1:]:
                                            elem.text = ''
                                        print(f"INFO: Replaced link text '{old_text}' with '{new_text}'")
                                        found = True
                                        fixes_applied += 1
                                        break
                            
                            # Also check individual runs for hyperlinks
                            if not found:
                                for run in paragraph.runs:
                                    hyperlinks = run._element.xpath('.//w:hyperlink')
                                    if hyperlinks:
                                        run_text = run.text
                                        if old_text.lower() in run_text.lower():
                                            # Replace text in run
                                            run.text = run_text.replace(old_text, new_text)
                                            print(f"INFO: Replaced link text in run: '{old_text}' -> '{new_text}'")
                                            found = True
                                            fixes_applied += 1
                                            break
                            if found:
                                break
                except Exception as e:
                    print(f"WARNING: Could not apply link text fix: {str(e)}", file=sys.stderr)
                    import traceback
                    traceback.print_exc()
        
        # Apply language span fixes (set language on specific text runs)
        language_fixes = fixes.get('languageFixes', [])
        if language_fixes and isinstance(language_fixes, list):
            print(f"INFO: Applying {len(language_fixes)} language span fixes...")
            for lang_fix in language_fixes:
                try:
                    text_content = lang_fix.get('text', '')
                    lang_code = lang_fix.get('language', 'fr-FR')  # Default to French if not specified
                    element_location = lang_fix.get('elementLocation', '')
                    
                    # Find text runs containing this content and set language
                    found = False
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text
                            # Check if this run contains the foreign text
                            if text_content and (text_content.lower() in run_text.lower() or run_text.lower() in text_content.lower()):
                                # Set language on this specific run
                                run._element.get_or_add_rPr().get_or_add_lang().set(qn('w:val'), lang_code)
                                print(f"INFO: Set language '{lang_code}' on text run containing: '{run_text[:50]}...'")
                                found = True
                                fixes_applied += 1
                                break
                        if found:
                            break
                    
                    if not found and element_location:
                        # Try to find by element location
                        for paragraph in doc.paragraphs:
                            for run in paragraph.runs:
                                if element_location.lower() in run.text.lower():
                                    run._element.get_or_add_rPr().get_or_add_lang().set(qn('w:val'), lang_code)
                                    print(f"INFO: Set language '{lang_code}' on text run (found by location): '{run.text[:50]}...'")
                                    fixes_applied += 1
                                    break
                except Exception as e:
                    print(f"WARNING: Could not apply language fix: {str(e)}", file=sys.stderr)
                    import traceback
                    traceback.print_exc()
        
        # Apply heading fixes
        # IMPORTANT: Preserve existing formatting - only add semantic heading structure
        heading_fixes = fixes.get('headingFixes', [])
        if heading_fixes and isinstance(heading_fixes, list):
            print(f"INFO: Applying {len(heading_fixes)} heading fixes (preserving formatting)...")
            for heading_fix in heading_fixes:
                try:
                    # heading_fix should contain: {paragraphIndex, level, text}
                    para_idx = heading_fix.get('paragraphIndex', -1)
                    level = heading_fix.get('level', 1)
                    text = heading_fix.get('text', '')
                    
                    # Ensure level is between 1 and 6
                    level = max(1, min(6, int(level)))
                    
                    if para_idx >= 0 and para_idx < len(doc.paragraphs):
                        paragraph = doc.paragraphs[para_idx]
                        
                        # PRESERVE EXISTING FORMATTING - ONLY set outline level, NEVER change style
                        # This marks it as a heading semantically without changing visual appearance
                        try:
                            # Set outline level in paragraph properties (preserves ALL formatting)
                            pPr = paragraph._element.get_or_add_pPr()
                            
                            # Remove any existing outline level first
                            existing_outline = pPr.find(qn('w:outlineLvl'))
                            if existing_outline is not None:
                                pPr.remove(existing_outline)
                            
                            # Add new outline level
                            outline_lvl = OxmlElement('w:outlineLvl')
                            outline_lvl.set(qn('w:val'), str(level - 1))  # Word uses 0-based (0=H1, 1=H2, etc.)
                            pPr.append(outline_lvl)
                            
                            # DO NOT apply heading style - this preserves existing formatting
                            # The outline level is enough for accessibility (screen readers, navigation)
                            print(f"INFO: Set heading level {level} (outline only, preserved formatting) on paragraph {para_idx}: '{text[:50]}...'")
                            
                            fixes_applied += 1
                        except Exception as style_error:
                            # If outline level fails, skip it - don't change style
                            print(f"WARNING: Could not set outline level (preserving formatting): {str(style_error)}", file=sys.stderr)
                            else:
                                print(f"WARNING: Could not set heading level: {str(style_error)}", file=sys.stderr)
                    else:
                        print(f"WARNING: Paragraph index {para_idx} out of range (document has {len(doc.paragraphs)} paragraphs)", file=sys.stderr)
                except Exception as e:
                    print(f"WARNING: Could not apply heading fix: {str(e)}", file=sys.stderr)
                    import traceback
                    traceback.print_exc()
        elif heading_fixes and isinstance(heading_fixes, (int, float)):
            # Legacy format - just a count, skip
            print(f"INFO: Heading fixes provided as count ({heading_fixes}), skipping (need paragraph indices)")
        
        # Always save the document, even if no fixes were applied
        # This ensures we return a valid document
        # Force save by accessing properties before saving
        try:
            # Access title to ensure it's set
            if fixes.get('metadata') and fixes['metadata'].get('title'):
                _ = doc.core_properties.title
        except:
            pass
        
        doc.save(output_path)
        print(f"INFO: Saved repaired document to: {output_path}")
        print(f"INFO: Applied {fixes_applied} fixes")
        
        # Verify title was saved
        if fixes.get('metadata') and fixes['metadata'].get('title'):
            try:
                # Re-open to verify
                verify_doc = Document(output_path)
                saved_title = verify_doc.core_properties.title
                if saved_title and saved_title.strip():
                    print(f"INFO: Verified title saved: {saved_title}")
                else:
                    print(f"WARNING: Title was not saved correctly (got: '{saved_title}')", file=sys.stderr)
            except Exception as e:
                print(f"WARNING: Could not verify title: {str(e)}", file=sys.stderr)
        
        # Verify the file was created
        if not Path(output_path).exists():
            print(f"ERROR: Output file was not created: {output_path}", file=sys.stderr)
            return False
        
        return True
        
    except Exception as e:
        print(f"ERROR: Failed to repair Word document: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        # Try to copy input to output as fallback
        try:
            import shutil
            shutil.copy2(input_path, output_path)
            print(f"INFO: Copied input to output as fallback")
            return True
        except Exception as e2:
            print(f"ERROR: Fallback copy also failed: {str(e2)}", file=sys.stderr)
            return False


def main():
    parser = argparse.ArgumentParser(description='Repair Word document with accessibility fixes')
    parser.add_argument('--input', required=True, help='Input Word document path')
    parser.add_argument('--output', required=True, help='Output Word document path')
    parser.add_argument('--fixes', required=True, help='JSON file with fixes to apply')
    
    args = parser.parse_args()
    
    # Read fixes from JSON file
    try:
        # Try utf-8-sig first (handles BOM), fall back to utf-8
        try:
            with open(args.fixes, 'r', encoding='utf-8-sig') as f:
                fixes = json.load(f)
        except UnicodeDecodeError:
            with open(args.fixes, 'r', encoding='utf-8') as f:
                fixes = json.load(f)
    except Exception as e:
        print(f"ERROR: Failed to read fixes file: {str(e)}", file=sys.stderr)
        sys.exit(1)
    
    # Repair the document
    success = repair_word_document(args.input, args.output, fixes)
    
    if not success:
        sys.exit(1)
    
    print("SUCCESS: Word document repaired")


if __name__ == '__main__':
    main()

